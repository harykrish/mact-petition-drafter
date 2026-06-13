"""All Opus 4.8 calls live here.

Four operations:
  extract_facts  — messy document  -> structured, sourced facts (JSON schema)
  kb_grade       — candidate KB     -> verdict vs kb_invariants.md (fresh context)
  draft_petition — knowledge base   -> MACT petition (KB facts only)
  petition_grade — petition + KB    -> verdict vs petition_rubric.md (fresh context)

Each verifier call is a *separate* messages.create with only the artifact +
rubric in context — there is no shared history with the drafter or extractor.
That separation is the "fresh context" the brief requires.
"""
from __future__ import annotations

import base64
import json
import os
import re
from pathlib import Path
from typing import Dict, List, Optional

from . import config

TEXT_EXTS = {".txt", ".md", ".json", ".csv"}
IMAGE_MEDIA = {".jpg": "image/jpeg", ".jpeg": "image/jpeg", ".png": "image/png",
               ".webp": "image/webp", ".gif": "image/gif"}
SUPPORTED_EXTS = set(TEXT_EXTS) | {".pdf", ".docx"} | set(IMAGE_MEDIA)

_CLIENT = None


def api_key_present() -> bool:
    return bool(os.environ.get("ANTHROPIC_API_KEY"))


def _client():
    global _CLIENT
    if _CLIENT is None:
        if not api_key_present():
            raise RuntimeError(
                "ANTHROPIC_API_KEY is not set. Set it in your environment (or .env / "
                "the Render/Railway dashboard) to run the LLM steps.")
        import anthropic
        import httpx
        # Fail fast and retry so a flaky/intermittent network can't wedge a run
        # (the SDK default is 10 min/request). Short connect timeout catches dead
        # connections; read timeout is per-chunk for streaming, so long drafts are
        # unaffected as long as tokens keep arriving.
        _CLIENT = anthropic.Anthropic(
            timeout=httpx.Timeout(60.0, connect=8.0), max_retries=4)
    return _CLIENT


def _text_of(message) -> str:
    return "".join(b.text for b in message.content if getattr(b, "type", None) == "text")


def _json_from_text(text: str) -> Dict:
    """Pull a JSON object out of a model response (last ```json fence, else outermost {...})."""
    # Greedy capture of fenced block contents (handles nested braces correctly).
    fences = re.findall(r"```(?:json)?\s*\n?(.*?)```", text, re.DOTALL)
    candidates = [f.strip() for f in fences if f.strip().startswith("{")]
    # Fallback: the outermost { ... } span.
    start, end = text.find("{"), text.rfind("}")
    if start != -1 and end > start:
        candidates.append(text[start:end + 1])
    for c in reversed(candidates):
        try:
            return json.loads(c)
        except json.JSONDecodeError:
            continue
    raise ValueError("No parseable JSON found in model response.")


def _reason_then_json(system: str, user: str, effort: str = "high",
                      max_tokens: int = 12000) -> Dict:
    """Adaptive-thinking call that ends in a fenced JSON block we parse.
    Streams so a long reasoning pass doesn't hit the non-streaming timeout guard."""
    client = _client()
    with client.messages.stream(
        model=config.MODEL,
        max_tokens=max_tokens,
        thinking={"type": "adaptive"},
        output_config={"effort": effort},
        system=system,
        messages=[{"role": "user", "content": user}],
    ) as stream:
        message = stream.get_final_message()
    return _json_from_text(_text_of(message))


# --- 1. extraction -------------------------------------------------------

_EXTRACT_SCHEMA = {
    "type": "object",
    "properties": {
        "facts": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "field": {"type": "string"},
                    "value": {"type": "string"},
                    "extraction_confidence": {"type": "number"},
                },
                "required": ["field", "value", "extraction_confidence"],
                "additionalProperties": False,
            },
        }
    },
    "required": ["facts"],
    "additionalProperties": False,
}


def _extract_system() -> str:
    field_lines = "\n".join("  - %s: %s" % (k, v) for k, v in config.CANONICAL_FIELDS.items())
    return (
        "You extract structured, sourced facts from a single legal/medical/financial "
        "document for an Indian Motor Accident Claims Tribunal (MACT) case.\n\n"
        "Rules:\n"
        "- Emit only facts STATED IN THIS DOCUMENT. Never infer or invent.\n"
        "- Use these canonical field names where applicable:\n" + field_lines + "\n"
        "- For each itemised hospital-bill line, emit a separate fact named "
        "'medical_expense_<slug>' (e.g. medical_expense_surgery) with the rupee amount. "
        "Do NOT emit a medical_expense_total — the line items are summed downstream.\n"
        "- Dates MUST be ISO YYYY-MM-DD. For a medical document, the accident_date is the "
        "date the accident occurred (often the admission date if it says the accident was "
        "that day), not the certificate/print date.\n"
        "- annual_income: integer rupees (gross annual). employment_type: exactly one of "
        "salaried_permanent / self_employed / fixed_wage.\n"
        "- value is always a string. extraction_confidence is 0..1.\n"
        "- Do not output provenance (stream/source) — the system attaches that."
    )


def extract_facts(doc_text: str, stream: str, source_doc: str, source_type: str) -> List[Dict]:
    client = _client()
    user = ("Document type: %s   Stream: %s   Source: %s\n\n--- DOCUMENT START ---\n%s\n--- DOCUMENT END ---"
            % (source_type, stream, source_doc, doc_text))
    message = client.messages.create(
        model=config.MODEL,
        max_tokens=4000,
        system=_extract_system(),
        messages=[{"role": "user", "content": user}],
        output_config={"format": {"type": "json_schema", "schema": _EXTRACT_SCHEMA}},
    )
    return _parse_facts(_json_from_text(_text_of(message)), stream, source_doc, source_type)


def _parse_facts(data: Dict, stream: str, source_doc: str, source_type: str) -> List[Dict]:
    raws = []
    for item in data.get("facts", []):
        field = str(item.get("field", "")).strip()
        if not field:
            continue
        raws.append({
            "field": field,
            "value": str(item.get("value", "")).strip(),
            "stream": stream,
            "source_doc": source_doc,
            "source_type": source_type,
            "extraction_confidence": float(item.get("extraction_confidence", 0.9)),
        })
    return raws


def _docx_text(path: str) -> str:
    from docx import Document  # python-docx
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            parts.append(" | ".join(c.text for c in row.cells))
    return "\n".join(parts)


# Schema for the vision path: documented facts PLUS optional AI visual observations.
_VISION_SCHEMA = {
    "type": "object",
    "properties": {
        "facts": _EXTRACT_SCHEMA["properties"]["facts"],
        "observations": {
            "type": "array",
            "items": {
                "type": "object",
                "properties": {
                    "finding": {"type": "string"},
                    "confidence": {"type": "number"},
                },
                "required": ["finding", "confidence"],
                "additionalProperties": False,
            },
        },
    },
    "required": ["facts"],
    "additionalProperties": False,
}

_VISION_GUIDANCE = (
    "\n\nThis input may be an image (a scan, photo, chart, or report page).\n"
    "Return TWO things:\n"
    "1) 'facts' — only information actually WRITTEN in the document (typed/printed/"
    "handwritten text, burned-in metadata: names, dates, study/modality, institution, "
    "printed values). These are documented facts.\n"
    "2) 'observations' — if this is a diagnostic image (CT/MRI/X-ray/ECG/ultrasound chart "
    "etc.), what you VISUALLY observe in the image itself (e.g. an apparent hyperdensity, "
    "fracture line, mass effect, abnormal trace). These are your interpretation, NOT "
    "documented facts — they must be confirmed by a clinician. If the input is plain typed "
    "text with no image to interpret, return an empty observations list.")


def _extract_with_blocks(blocks: List[Dict], stream: str, source_doc: str, source_type: str,
                         extract_observations: bool = True) -> List[Dict]:
    """Extraction from non-text input (PDF/image content blocks) via Opus vision.
    Returns documented facts (under `source_type`) plus, optionally, AI visual
    observations recorded as low-confidence, review-flagged image_finding facts."""
    client = _client()
    instruction = ("Document type: %s   Stream: %s   Source: %s\nExtract the structured facts "
                   "from the attached document per the rules." % (source_type, stream, source_doc))
    schema = _EXTRACT_SCHEMA
    if extract_observations:
        instruction += _VISION_GUIDANCE
        schema = _VISION_SCHEMA
    content = list(blocks) + [{"type": "text", "text": instruction}]
    message = client.messages.create(
        model=config.MODEL,
        max_tokens=4000,
        system=_extract_system(),
        messages=[{"role": "user", "content": content}],
        output_config={"format": {"type": "json_schema", "schema": schema}},
    )
    data = _json_from_text(_text_of(message))
    raws = _parse_facts(data, stream, source_doc, source_type)
    for obs in data.get("observations", []) or []:
        finding = str(obs.get("finding", "")).strip()
        if finding:
            raws.append({
                "field": "image_finding",
                "value": finding,
                "stream": stream,
                "source_doc": source_doc,
                "source_type": "AI Vision Observation",
                "extraction_confidence": float(obs.get("confidence", 0.4)),
            })
    return raws


def extract_facts_from_path(path: str, stream: str, source_doc: str, source_type: str,
                            extract_observations: bool = True) -> List[Dict]:
    """Multi-format extraction: text/markdown/json/csv, .docx, .pdf, and images.
    For image/PDF inputs, also captures AI visual observations (flagged, low-confidence)."""
    ext = Path(path).suffix.lower()
    if ext in TEXT_EXTS:
        return extract_facts(Path(path).read_text(encoding="utf-8", errors="ignore"),
                             stream, source_doc, source_type)
    if ext == ".docx":
        return extract_facts(_docx_text(path), stream, source_doc, source_type)
    if ext == ".pdf":
        data = base64.standard_b64encode(Path(path).read_bytes()).decode("ascii")
        block = {"type": "document", "source": {"type": "base64",
                 "media_type": "application/pdf", "data": data}}
    elif ext in IMAGE_MEDIA:
        block = _image_block(path)   # vision; auto-downscaled to fit API limits
    else:
        raise ValueError("unsupported file type: %s" % ext)
    return _extract_with_blocks([block], stream, source_doc, source_type,
                                extract_observations=extract_observations)


def _image_block(path: str, max_edge: int = 2200, max_bytes: int = 4_500_000) -> Dict:
    """Read a scan/photo as an Opus vision block. Large images are downscaled and
    re-encoded so they fit the API's per-image size limit (never skipped)."""
    import io
    try:
        from PIL import Image
    except Exception:
        data = base64.standard_b64encode(Path(path).read_bytes()).decode("ascii")
        mt = IMAGE_MEDIA.get(Path(path).suffix.lower(), "image/jpeg")
        return {"type": "image", "source": {"type": "base64", "media_type": mt, "data": data}}
    img = Image.open(path)
    if img.mode not in ("RGB", "L"):
        img = img.convert("RGB")
    if max(img.size) > max_edge:
        r = max_edge / float(max(img.size))
        img = img.resize((max(1, int(img.size[0] * r)), max(1, int(img.size[1] * r))))
    data = b""
    for q in (85, 70, 55, 40, 30):
        buf = io.BytesIO()
        img.convert("RGB").save(buf, format="JPEG", quality=q)
        data = buf.getvalue()
        if len(data) <= max_bytes:
            break
    return {"type": "image", "source": {"type": "base64", "media_type": "image/jpeg",
            "data": base64.standard_b64encode(data).decode("ascii")}}


# --- 2. KB grading (fresh context) --------------------------------------

def kb_grade(record: Dict) -> Dict:
    rubric = Path(config.KB_INVARIANTS_PATH).read_text(encoding="utf-8")
    system = (
        "You are an INDEPENDENT knowledge-base verifier. You have no memory of how this "
        "record was built. Grade the candidate case_record.json strictly against the "
        "invariants. Reason about each MUST, then end your reply with a single fenced "
        "```json block matching the rubric's output contract "
        "(keys: result, must[], should[], blocking_failures[]).")
    user = ("# kb_invariants.md\n\n%s\n\n# candidate case_record.json\n\n```json\n%s\n```"
            % (rubric, json.dumps(record, indent=2, ensure_ascii=False)))
    return _reason_then_json(system, user, effort="medium", max_tokens=8000)


# --- 3. drafting ---------------------------------------------------------

def draft_petition(record: Dict, feedback: Optional[str] = None) -> str:
    from . import store
    active = store.active_facts(record)
    facts_view = [{"id": f["id"], "field": f["field"], "value": f["value"],
                   "stream": f["stream"], "source_type": f["source_type"],
                   "needs_human_review": f.get("needs_human_review", False)} for f in active]
    contradictions = record.get("contradictions", [])
    precedent = "\n\n".join(
        Path(p).read_text(encoding="utf-8")
        for p in sorted(Path(config.PRECEDENT_DIR).glob("*.md")))

    system = (
        "You are a claimant's advocate drafting a compensation petition before a Motor "
        "Accident Claims Tribunal (MACT) in India.\n\n"
        "HARD RULES:\n"
        "- Use ONLY the facts in the knowledge base provided. Every factual assertion "
        "(name, age, date, income, disability %, expense, liability) MUST cite the fact "
        "id(s) it rests on, e.g. [F003].\n"
        "- Liability/negligence must come from police-stream facts; income from financial; "
        "disability from medical.\n"
        "- Any fact with needs_human_review=true, or any value listed in contradictions, "
        "MUST NOT be stated as settled — flag it as disputed (note both values and which "
        "you adopt and why).\n"
        "- Facts with source_type 'AI Vision Observation' are AI interpretations of a scan/"
        "image, NOT documented evidence. Never cite them as established findings; rely on the "
        "written radiology reports instead. At most note that they warrant radiologist "
        "confirmation.\n"
        "- Compute the loss of future earning capacity as: annual_income x (1 + future "
        "prospects) x multiplier x functional_disability%, using the Sarla Verma multiplier "
        "table and the Pranay Sethi future-prospects table from the precedent notes. SHOW "
        "the multiplier (with age band) and the future-prospects fraction (with band).\n"
        "- Medical expenses = sum of the itemised medical_expense_* facts.\n"
        "- Cite Sarla Verma (2009) and Pranay Sethi (2017); reference the 2024 authority "
        "note only as reinforcing (its citation is unconfirmed — do not overstate it).\n"
        "- This is a legal document of documented facts. No medical advice.\n\n"
        "STRUCTURE (markdown):\n"
        "1. Cause title (the Tribunal; claimant vs respondents: driver/owner/insurer).\n"
        "2. Parties.\n3. Facts of the accident (with liability).\n4. Injuries & permanent "
        "disability.\n5. Income & avocation.\n6. Heads of compensation — include a markdown "
        "table titled 'SCHEDULE OF COMPENSATION' with columns: Head | Amount (INR) | Basis. "
        "One row per head, a final **TOTAL** row. Amounts are plain integers (no commas in "
        "the Amount column).\n7. Prayer.\n8. Precedents relied upon.")
    user = ("# KNOWLEDGE BASE — active facts\n```json\n%s\n```\n\n"
            "# CONTRADICTIONS (unresolved — flag, do not assert)\n```json\n%s\n```\n\n"
            "# PRECEDENT NOTES\n%s\n"
            % (json.dumps(facts_view, indent=2, ensure_ascii=False),
               json.dumps(contradictions, indent=2, ensure_ascii=False),
               precedent))
    if feedback:
        user += ("\n\n# VERIFIER FEEDBACK on your previous draft — fix every item:\n%s" % feedback)

    client = _client()
    with client.messages.stream(
        model=config.MODEL,
        max_tokens=16000,
        thinking={"type": "adaptive"},
        output_config={"effort": "high"},
        system=system,
        messages=[{"role": "user", "content": user}],
    ) as stream:
        message = stream.get_final_message()
    return _text_of(message)


# --- 4. petition grading (fresh context) --------------------------------

def petition_grade(petition_md: str, record: Dict) -> Dict:
    from . import store
    rubric = Path(config.PETITION_RUBRIC_PATH).read_text(encoding="utf-8")
    active = store.active_facts(record)
    facts_view = [{"id": f["id"], "field": f["field"], "value": f["value"],
                   "stream": f["stream"]} for f in active]
    system = (
        "You are an INDEPENDENT petition verifier. You did not write this petition and have "
        "no memory of drafting it. Grade it strictly against the rubric, checking that every "
        "factual assertion traces to a knowledge-base fact id, that disputed facts are not "
        "asserted as settled, and that liability/income/disability come from the right "
        "streams. Independently PARSE every amount from the SCHEDULE OF COMPENSATION. Reason "
        "through each MUST, then end with a single fenced ```json block matching the rubric's "
        "output contract, and additionally include a 'petition_claimed' object with the "
        "amounts you parsed: {loss_of_earning, medical_expenses, heads:{<head>:<amount>}, "
        "grand_total}.")
    user = ("# petition_rubric.md\n\n%s\n\n# KNOWLEDGE BASE — active facts\n```json\n%s\n```\n\n"
            "# PETITION DRAFT\n\n%s" % (rubric, json.dumps(facts_view, indent=2, ensure_ascii=False), petition_md))
    return _reason_then_json(system, user, effort="high", max_tokens=12000)
